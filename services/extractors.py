"""
Извлечение текста из файлов разных форматов: .txt, .docx, .pdf
Каждая функция возвращает "сырой" текст без форматирования
"""
import logging

logger = logging.getLogger(__name__)


def extract_from_txt(path: str) -> str:
    """Читает обычный текстовый файл, пробуя разные кодировки"""
    encodings = ["utf-8", "cp1251", "latin-1"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    # если ничего не подошло — читаем с игнорированием ошибок
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_from_docx(path: str) -> str:
    """Извлекает текст из .docx (абзацы + таблицы)"""
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                parts.append(" | ".join(cells_text))

    return "\n".join(parts)


def extract_from_pdf(path: str) -> str:
    """Извлекает текст из .pdf постранично через pdfplumber"""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)

    text = "\n".join(text_parts)

    if not text.strip():
        logger.warning(
            "PDF не содержит извлекаемого текстового слоя (похоже на скан). "
            "OCR в этом боте не реализован."
        )
    return text


EXTRACTORS = {
    ".txt": extract_from_txt,
    ".docx": extract_from_docx,
    ".pdf": extract_from_pdf,
}


def extract_text(path: str, extension: str) -> str:
    """
    Диспетчер: выбирает нужную функцию извлечения по расширению файла.
    extension должен быть в формате ".txt", ".docx", ".pdf" (в нижнем регистре)
    """
    extension = extension.lower()
    if extension not in EXTRACTORS:
        raise ValueError(f"Формат {extension} не поддерживается")
    return EXTRACTORS[extension](path)
